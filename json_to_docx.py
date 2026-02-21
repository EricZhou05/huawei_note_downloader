# -*- coding: utf-8 -*-
import json
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, font_name, size, bold=False, italic=False):
    """设置字体、字号、加粗、斜体以及中文字体支持"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_horizontal_line(doc):
    """在文档中添加一条明显的水平分割线（3pt粗黑色实线）"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '24')  # 24 = 3pt，非常明显
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def clean_text_for_comparison(text):
    """仅保留汉字，用于忽略格式、标点、数字、英文的干扰"""
    if not text:
        return ""
    # 使用正则仅提取汉字
    hanzi_list = re.findall(r'[\u4e00-\u9fa5]', text)
    if not hanzi_list:
        # 如果没有汉字（比如纯英文笔记），回退到去除空白字符
        return "".join(text.split())
    return "".join(hanzi_list)

def json_to_docx(json_path, output_path):
    if not os.path.exists(json_path):
        print(f"错误: 找不到文件 {json_path}")
        return

    with open(json_path, 'r', encoding='utf-8') as f:
        notes = json.load(f)

    doc = Document()

    # --- 样式预处理：修改内置样式的基准字体 ---
    # 即使应用了样式，Word有时会回退到样式默认字体。
    # 这里修改文档级别的 Heading 1 和 Normal 定义，作为双重保险。
    try:
        style_h1 = doc.styles['Heading 1']
        style_h1.font.name = '微软雅黑'
        style_h1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        style_h1.font.color.rgb = RGBColor(0, 0, 0)
        
        style_normal = doc.styles['Normal']
        style_normal.font.name = '宋体'
        style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style_normal.font.color.rgb = RGBColor(0, 0, 0)
    except Exception as e:
        print(f"警告: 样式定义修改失败，将依赖手动覆盖。错误: {e}")

    for index, note in enumerate(notes):
        content_raw = note.get('content', '')
        title_raw = note.get('title') or '无标题'
        
        # 1. 标题处理
        # 预处理：去除换行符，便于显示
        title = title_raw.replace('\n', ' ').replace('\r', ' ').strip()
        
        # 智能缩减逻辑：仅比较汉字
        cmp_title = clean_text_for_comparison(title)
        cmp_content = clean_text_for_comparison(content_raw)
        
        # 仅处理有一定长度的标题，防止误伤短标题（如"超市"、"书单"）
        if len(title) > 6 and len(cmp_content) > 0:
            # 策略升级：反向包含检测
            # 检查正文的开头（前10个汉字），是否包含在标题的开头（前50个汉字）中
            # 这种方式能覆盖 "Title = Content" 以及 "Title = 摘要 + Content" 的情况
            
            # 取正文前10个汉字作为特征指纹
            content_fingerprint = cmp_content[:10]
            # 取标题前50个汉字作为搜索区域
            title_search_area = cmp_title[:50]
            
            # 只要特征指纹够长（避免误判），且出现在标题里
            if len(content_fingerprint) >= 4 and content_fingerprint in title_search_area:
                # 判定为自动生成的冗余标题，截取前12个原始字符
                title = title[:12]
        
        # --- 核心：应用 Heading 1 样式 ---
        # 使用 style='Heading 1' 让 Word 识别这是标题，从而能生成目录
        title_para = doc.add_paragraph(style='Heading 1')
        title_run = title_para.add_run(title)
        # 再次强制覆盖字体，确保是你要求的格式而不是 Word 默认的蓝色标题
        set_font(title_run, '微软雅黑', 16, bold=True)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 2. 时间信息
        created_time = note.get('created_time', '未知')
        modified_time = note.get('modified_time', '未知')
        time_info = f"创建时间: {created_time}  |  修改时间: {modified_time}"
        
        time_para = doc.add_paragraph(style='Normal')
        time_run = time_para.add_run(time_info)
        set_font(time_run, '微软雅黑', 9)
        time_para.paragraph_format.space_after = Pt(6) # 增加一点段后距
        
        # 3. 正文内容
        if content_raw:
            # 压缩连续换行
            content = re.sub(r'\n+', '\n', content_raw).strip()
            
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    # 应用 Normal 样式
                    content_para = doc.add_paragraph(style='Normal')
                    content_run = content_para.add_run(line)
                    set_font(content_run, '宋体', 11)
                else:
                    doc.add_paragraph(style='Normal')
        else:
            content_para = doc.add_paragraph(style='Normal')
            content_run = content_para.add_run("[无内容]")
            set_font(content_run, '宋体', 11, italic=True)

        # 4. 分割线
        if index < len(notes) - 1:
            doc.add_paragraph(style='Normal') # 上空行
            add_horizontal_line(doc)          # 粗实线
            doc.add_paragraph(style='Normal') # 下空行

    # 保存文档
    doc.save(output_path)
    print(f"转换完成！优化后的 Word 文档已保存至: {output_path}")

if __name__ == "__main__":
    current_dir = os.path.dirname(os.path.abspath(__file__))
    input_json = os.path.join(current_dir, "华为备忘录导出.json")
    output_docx = os.path.join(current_dir, "华为备忘录导出.docx")
    
    json_to_docx(input_json, output_docx)
